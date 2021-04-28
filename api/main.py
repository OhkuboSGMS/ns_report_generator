from docx import Document
from pprint import pprint
file_name = "./test_data/report.docx"

document = Document(file_name)

for i,par in enumerate(document.paragraphs):
    print(i,par.text)

for i,tbl in enumerate(document.tables):
    print(f"Table:{i}")
    for i_row,row in enumerate(tbl.rows):
        print(f"{i_row}:",end="")
        for cell in row.cells:
                # if len(cell.paragraphs[0].runs) >0:
                #     print("Runs",cell.paragraphs[0].runs[0].text)
                print("|"+cell.text.replace("\n","")[:3],end="")
        print("--------------")

tbl = document.tables[0]
#TODO スタイルをコピーする
# apply date
pprint(tbl.rows[4].cells[0])
print("Before,")

for i,run in enumerate(tbl.rows[4].cells[0].paragraphs[0].runs):
    print(i,run.text)

tbl.rows[4].cells[0].paragraphs[0].runs[2].text ="2020年　10月 24日"
pprint(tbl.rows[4].cells[0].paragraphs[0].runs[0].text)
# department
tbl.rows[6].cells[0].paragraphs[0].runs[0].text="所属：　システム開発部　氏名：　大久保成将　㊞"
#type
tbl.rows[10].cells[1].paragraphs[0].runs[0].text =" yuukyuu"
# duration
# date
# reason
# address , TEL
# etc 振替、代休

print(f"Depart {tbl.rows[6].cells[0].text}")

document.save("test.docx")

